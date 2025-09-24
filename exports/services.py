"""
Export Services for DidactIA

This module provides services for exporting educational content to various formats
including PDF, DOCX, HTML, and ZIP archives.
"""

import io
import os
import json
import zipfile
from datetime import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path

from django.conf import settings
from django.template.loader import render_to_string
from django.utils import timezone
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage

# Import export libraries
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

import logging

logger = logging.getLogger(__name__)


class PDFExporter:
    """Service for exporting content to PDF format using ReportLab"""
    
    def __init__(self):
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab is required for PDF export")
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom paragraph styles"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        ))
        
        # Header style
        self.styles.add(ParagraphStyle(
            name='CustomHeader',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=12,
            textColor=colors.darkblue,
            borderWidth=1,
            borderColor=colors.darkblue,
            borderPadding=5
        ))
        
        # Question style
        self.styles.add(ParagraphStyle(
            name='Question',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceBefore=10,
            spaceAfter=5,
            leftIndent=20
        ))
        
        # Answer style
        self.styles.add(ParagraphStyle(
            name='Answer',
            parent=self.styles['Normal'],
            fontSize=10,
            leftIndent=40,
            textColor=colors.darkgreen
        ))
    
    def export_quiz(self, quiz_data: Dict[str, Any], branding: Dict[str, Any] = None) -> io.BytesIO:
        """Export quiz to PDF format"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Add header/branding
        if branding:
            story.extend(self._add_branding(branding))
        
        # Title
        title = quiz_data.get('title', 'Quiz')
        story.append(Paragraph(title, self.styles['CustomTitle']))
        story.append(Spacer(1, 20))
        
        # Description
        if quiz_data.get('description'):
            story.append(Paragraph(quiz_data['description'], self.styles['Normal']))
            story.append(Spacer(1, 15))
        
        # Instructions
        instructions = f"""
        <b>Instructions:</b><br/>
        • Read each question carefully<br/>
        • Choose the best answer for multiple choice questions<br/>
        • Duration: {quiz_data.get('estimated_duration', 'Not specified')}<br/>
        • Total Points: {quiz_data.get('total_points', len(quiz_data.get('questions', [])))}
        """
        story.append(Paragraph(instructions, self.styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Questions
        questions = quiz_data.get('questions', [])
        for i, question in enumerate(questions, 1):
            question_type = question.get('type', 'multiple_choice')
            question_type_display = {
                'multiple_choice': 'Multiple Choice',
                'true_false': 'True/False',
                'short_answer': 'Short Answer',
                'fill_blank': 'Fill in the Blank',
                'essay': 'Essay'
            }.get(question_type, 'Multiple Choice')
            
            # Question header with type
            type_text = f"<b>[{question_type_display}]</b>"
            story.append(Paragraph(type_text, self.styles['Answer']))
            
            # Question text
            q_text = f"<b>{i}. {question.get('question', '')}</b> ({question.get('points', 1)} point{'s' if question.get('points', 1) != 1 else ''})"
            story.append(Paragraph(q_text, self.styles['Question']))
            
            # Handle different question types
            if question_type == 'multiple_choice' and question.get('options'):
                for j, option in enumerate(question['options']):
                    option_letter = chr(65 + j)  # A, B, C, D
                    story.append(Paragraph(f"☐ {option_letter}. {option}", self.styles['Normal']))
            
            elif question_type == 'true_false':
                story.append(Paragraph("☐ True", self.styles['Normal']))
                story.append(Paragraph("☐ False", self.styles['Normal']))
            
            elif question_type == 'short_answer':
                story.append(Spacer(1, 10))
                story.append(Paragraph("Answer: " + "_" * 70, self.styles['Normal']))
            
            elif question_type == 'fill_blank':
                story.append(Spacer(1, 10))
                story.append(Paragraph("Answer: " + "_" * 70, self.styles['Normal']))
            
            elif question_type == 'essay':
                story.append(Spacer(1, 10))
                for _ in range(5):
                    story.append(Paragraph("_" * 100, self.styles['Normal']))
            
            story.append(Spacer(1, 15))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_exam(self, exam_data: Dict[str, Any], branding: Dict[str, Any] = None) -> io.BytesIO:
        """Export exam to PDF format"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Add header/branding
        if branding:
            story.extend(self._add_branding(branding))
        
        # Title
        title = exam_data.get('title', 'Comprehensive Exam')
        story.append(Paragraph(title, self.styles['CustomTitle']))
        story.append(Spacer(1, 20))
        
        # Exam info
        info_data = [
            ['Duration:', f"{exam_data.get('duration', 120)} minutes"],
            ['Total Questions:', str(exam_data.get('total_questions', 0))],
            ['Date:', datetime.now().strftime('%B %d, %Y')],
            ['Time:', '________________'],
            ['Student Name:', '_' * 30],
            ['Student ID:', '_' * 20]
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 3*inch])
        info_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # Instructions
        story.append(Paragraph("<b>INSTRUCTIONS:</b>", self.styles['CustomHeader']))
        instructions = """
        1. Read all instructions carefully before beginning<br/>
        2. Write your answers clearly in the spaces provided<br/>
        3. For multiple choice questions, circle the letter of your answer<br/>
        4. Show all work for calculation problems<br/>
        5. Check your answers before submitting
        """
        story.append(Paragraph(instructions, self.styles['Normal']))
        story.append(PageBreak())
        
        # Exam sections
        question_num = 1
        for section in exam_data.get('sections', []):
            # Section header
            section_title = f"Section {len(story)//20 + 1}: {section.get('name', 'Questions')}"
            story.append(Paragraph(section_title, self.styles['CustomHeader']))
            
            if section.get('instructions'):
                story.append(Paragraph(section['instructions'], self.styles['Normal']))
                story.append(Spacer(1, 10))
            
            # Section questions
            for question in section.get('questions', []):
                # Question text
                q_text = f"<b>{question_num}. {question.get('question', '')}</b> ({question.get('points', 1)} points)"
                story.append(Paragraph(q_text, self.styles['Question']))
                
                # Question options or answer space
                if question.get('type') == 'multiple_choice' and question.get('options'):
                    for j, option in enumerate(question['options']):
                        option_letter = chr(65 + j)
                        story.append(Paragraph(f"{option_letter}. {option}", self.styles['Normal']))
                else:
                    # Add answer space
                    for _ in range(3):  # 3 lines for answer
                        story.append(Paragraph("_" * 70, self.styles['Normal']))
                        story.append(Spacer(1, 5))
                
                story.append(Spacer(1, 15))
                question_num += 1
            
            story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_answer_key(self, content_data: Dict[str, Any], branding: Dict[str, Any] = None) -> io.BytesIO:
        """Export answer key to PDF format"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Add header/branding
        if branding:
            story.extend(self._add_branding(branding))
        
        # Title
        title = f"{content_data.get('title', 'Quiz')} - Answer Key"
        story.append(Paragraph(title, self.styles['CustomTitle']))
        story.append(Spacer(1, 20))
        
        # Answers
        if 'questions' in content_data:
            questions = content_data['questions']
        else:
            # Handle exam format with sections
            questions = []
            for section in content_data.get('sections', []):
                questions.extend(section.get('questions', []))
        
        for i, question in enumerate(questions, 1):
            # Question and answer
            q_text = f"<b>{i}. {question.get('question', '')}</b>"
            story.append(Paragraph(q_text, self.styles['Question']))
            
            answer_text = f"<b>Answer:</b> {question.get('correct_answer', 'Not provided')}"
            story.append(Paragraph(answer_text, self.styles['Answer']))
            
            # Explanation if available
            if question.get('explanation'):
                exp_text = f"<b>Explanation:</b> {question['explanation']}"
                story.append(Paragraph(exp_text, self.styles['Normal']))
            
            story.append(Spacer(1, 10))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def _add_branding(self, branding: Dict[str, Any]) -> List:
        """Add branding elements to document"""
        elements = []
        
        # Institution name
        if branding.get('institution_name'):
            inst_style = ParagraphStyle(
                name='Institution',
                parent=self.styles['Normal'],
                fontSize=14,
                alignment=TA_CENTER,
                spaceBefore=0,
                spaceAfter=10,
                textColor=colors.darkblue
            )
            elements.append(Paragraph(branding['institution_name'], inst_style))
        
        # Department
        if branding.get('department'):
            dept_style = ParagraphStyle(
                name='Department',
                parent=self.styles['Normal'],
                fontSize=12,
                alignment=TA_CENTER,
                spaceAfter=20
            )
            elements.append(Paragraph(branding['department'], dept_style))
        
        return elements


class DOCXExporter:
    """Service for exporting content to DOCX format"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export")
    
    def export_quiz(self, quiz_data: Dict[str, Any], branding: Dict[str, Any] = None) -> io.BytesIO:
        """Export quiz to DOCX format"""
        doc = Document()
        
        # Add branding
        if branding:
            self._add_branding(doc, branding)
        
        # Title
        title = doc.add_heading(quiz_data.get('title', 'Quiz'), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Description
        if quiz_data.get('description'):
            doc.add_paragraph(quiz_data['description'])
        
        # Instructions
        doc.add_heading('Instructions:', level=2)
        instructions = doc.add_paragraph()
        instructions.add_run('• Read each question carefully\n')
        instructions.add_run('• Choose the best answer for multiple choice questions\n')
        instructions.add_run(f'• Duration: {quiz_data.get("estimated_duration", "Not specified")}\n')
        instructions.add_run(f'• Total Points: {quiz_data.get("total_points", len(quiz_data.get("questions", [])))}')
        
        # Questions
        for i, question in enumerate(quiz_data.get('questions', []), 1):
            question_type = question.get('type', 'multiple_choice')
            question_type_display = {
                'multiple_choice': 'Multiple Choice',
                'true_false': 'True/False',
                'short_answer': 'Short Answer',
                'fill_blank': 'Fill in the Blank',
                'essay': 'Essay'
            }.get(question_type, 'Multiple Choice')
            
            # Question type header
            type_para = doc.add_paragraph()
            type_para.add_run(f'[{question_type_display}]').bold = True
            
            # Question text
            q_para = doc.add_paragraph()
            q_para.add_run(f'{i}. {question.get("question", "")} ').bold = True
            q_para.add_run(f'({question.get("points", 1)} point{"s" if question.get("points", 1) != 1 else ""})')
            
            # Handle different question types
            if question_type == 'multiple_choice' and question.get('options'):
                for j, option in enumerate(question['options']):
                    option_letter = chr(65 + j)
                    doc.add_paragraph(f'☐ {option_letter}. {option}')
            
            elif question_type == 'true_false':
                doc.add_paragraph('☐ True')
                doc.add_paragraph('☐ False')
            
            elif question_type == 'short_answer':
                doc.add_paragraph('Answer: ' + '_' * 70)
            
            elif question_type == 'fill_blank':
                doc.add_paragraph('Answer: ' + '_' * 70)
            
            elif question_type == 'essay':
                for _ in range(5):
                    doc.add_paragraph('_' * 80)
            
            doc.add_paragraph()  # Add space between questions
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def export_exam(self, exam_data: Dict[str, Any], branding: Dict[str, Any] = None) -> io.BytesIO:
        """Export exam to DOCX format"""
        doc = Document()
        
        # Add branding
        if branding:
            self._add_branding(doc, branding)
        
        # Title
        title = doc.add_heading(exam_data.get('title', 'Comprehensive Exam'), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Exam info table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        info_data = [
            ('Duration:', f"{exam_data.get('duration', 120)} minutes"),
            ('Total Questions:', str(exam_data.get('total_questions', 0))),
            ('Date:', datetime.now().strftime('%B %d, %Y')),
            ('Time:', '________________'),
            ('Student Name:', '_' * 30),
            ('Student ID:', '_' * 20)
        ]
        
        for i, (label, value) in enumerate(info_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph()
        
        # Instructions
        doc.add_heading('INSTRUCTIONS:', level=2)
        instructions = doc.add_paragraph()
        instructions.add_run('1. Read all instructions carefully before beginning\n')
        instructions.add_run('2. Write your answers clearly in the spaces provided\n')
        instructions.add_run('3. For multiple choice questions, circle the letter of your answer\n')
        instructions.add_run('4. Show all work for calculation problems\n')
        instructions.add_run('5. Check your answers before submitting')
        
        doc.add_page_break()
        
        # Exam sections
        question_num = 1
        for section in exam_data.get('sections', []):
            # Section header
            section_title = f"Section: {section.get('name', 'Questions')}"
            doc.add_heading(section_title, level=2)
            
            if section.get('instructions'):
                doc.add_paragraph(section['instructions'])
            
            # Section questions
            for question in section.get('questions', []):
                # Question text
                q_para = doc.add_paragraph()
                q_para.add_run(f'{question_num}. {question.get("question", "")} ').bold = True
                q_para.add_run(f'({question.get("points", 1)} points)')
                
                # Options or answer space
                if question.get('type') == 'multiple_choice' and question.get('options'):
                    for j, option in enumerate(question['options']):
                        option_letter = chr(65 + j)
                        doc.add_paragraph(f'{option_letter}. {option}', style='List Number')
                else:
                    # Add answer lines
                    for _ in range(3):
                        doc.add_paragraph('_' * 70)
                
                doc.add_paragraph()
                question_num += 1
            
            doc.add_page_break()
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _add_branding(self, doc: Document, branding: Dict[str, Any]):
        """Add branding to document"""
        if branding.get('institution_name'):
            inst_para = doc.add_paragraph(branding['institution_name'])
            inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inst_para.runs[0].font.size = Pt(14)
            inst_para.runs[0].bold = True
        
        if branding.get('department'):
            dept_para = doc.add_paragraph(branding['department'])
            dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dept_para.runs[0].font.size = Pt(12)
        
        doc.add_paragraph()


class HTMLExporter:
    """Service for exporting content to HTML format with professional university styling"""
    
    def export_quiz(self, quiz_data: Dict[str, Any], branding: Dict[str, Any] = None, show_answers: bool = False) -> str:
        """Export quiz to HTML format
        
        Args:
            quiz_data: The quiz data to export
            branding: Optional branding information
            show_answers: Whether to show correct answers (for instructor version)
        """
        context = {
            'quiz_data': quiz_data,
            'branding': branding or {},
            'export_date': datetime.now().strftime('%B %d, %Y')
        }
        
        html_template = """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{{ quiz_data.title }}</title>
            <style>
                @import url('https://fonts.googleapis.com/css2?family=Crimson+Text:wght@400;600;700&family=Inter:wght@300;400;500;600&display=swap');
                
                * {
                    margin: 0;
                    padding: 0;
                    box-sizing: border-box;
                }
                
                body {
                    font-family: 'Inter', 'Arial', sans-serif;
                    line-height: 1.6;
                    color: #2c3e50;
                    background: white;
                    margin: 0;
                    padding: 60px 40px 40px;
                }
                
                .exam-container {
                    max-width: 210mm;
                    margin: 0 auto;
                    background: white;
                    min-height: 297mm;
                }
                
                /* University Header */
                .university-header {
                    display: flex;
                    align-items: center;
                    border-bottom: 3px solid #2c3e50;
                    padding-bottom: 20px;
                    margin-bottom: 30px;
                }
                
                .logo-placeholder {
                    width: 80px;
                    height: 80px;
                    border: 2px solid #34495e;
                    border-radius: 50%;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    font-size: 12px;
                    color: #7f8c8d;
                    margin-right: 25px;
                    background: #ecf0f1;
                    flex-shrink: 0;
                }
                
                .header-info {
                    flex: 1;
                }
                
                .university-name {
                    font-family: 'Crimson Text', serif;
                    font-size: 28px;
                    font-weight: 700;
                    color: #2c3e50;
                    margin-bottom: 5px;
                    text-transform: uppercase;
                    letter-spacing: 1px;
                }
                
                .department-name {
                    font-size: 16px;
                    color: #34495e;
                    margin-bottom: 4px;
                    font-weight: 500;
                }
                
                .course-info {
                    font-size: 14px;
                    color: #7f8c8d;
                    font-weight: 400;
                }
                
                /* Exam Title */
                .exam-title {
                    text-align: center;
                    margin: 40px 0 30px;
                }
                
                .exam-type {
                    font-family: 'Crimson Text', serif;
                    font-size: 36px;
                    font-weight: 700;
                    color: #2c3e50;
                    margin-bottom: 10px;
                    text-transform: uppercase;
                    letter-spacing: 2px;
                }
                
                /* Exam Metadata */
                .exam-metadata {
                    display: grid;
                    grid-template-columns: 1fr 1fr;
                    gap: 15px;
                    margin: 30px 0;
                    padding: 20px;
                    border: 1px solid #bdc3c7;
                    background: #f8f9fa;
                }
                
                .metadata-item {
                    display: flex;
                    align-items: center;
                    font-size: 14px;
                }
                
                .metadata-label {
                    font-weight: 600;
                    color: #2c3e50;
                    margin-right: 10px;
                    min-width: 80px;
                }
                
                .metadata-value {
                    color: #34495e;
                    border-bottom: 1px solid #bdc3c7;
                    flex: 1;
                    padding-bottom: 2px;
                }
                
                /* Instructions Section */
                .instructions-section {
                    margin: 30px 0;
                    padding: 20px;
                    border: 2px solid #34495e;
                    background: #fdfdfd;
                }
                
                .instructions-title {
                    font-family: 'Crimson Text', serif;
                    font-size: 18px;
                    font-weight: 700;
                    color: #2c3e50;
                    margin-bottom: 15px;
                    text-transform: uppercase;
                    border-bottom: 1px solid #ecf0f1;
                    padding-bottom: 5px;
                }
                
                .instructions-list {
                    list-style: none;
                }
                
                .instructions-list li {
                    padding: 5px 0;
                    position: relative;
                    padding-left: 20px;
                    color: #34495e;
                    font-size: 14px;
                }
                
                .instructions-list li::before {
                    content: '•';
                    position: absolute;
                    left: 0;
                    color: #2c3e50;
                    font-weight: bold;
                }
                
                /* Questions */
                .question {
                    margin: 25px 0;
                    page-break-inside: avoid;
                }
                
                .question-header {
                    display: flex;
                    justify-content: space-between;
                    align-items: flex-start;
                    margin-bottom: 12px;
                }
                
                .question-number {
                    font-weight: 600;
                    color: #2c3e50;
                    font-size: 16px;
                }
                
                .question-points {
                    font-size: 12px;
                    color: #7f8c8d;
                    font-style: italic;
                }
                
                .question-text {
                    font-size: 15px;
                    line-height: 1.6;
                    color: #2c3e50;
                    margin-bottom: 15px;
                    text-align: justify;
                }
                
                /* Multiple Choice Options */
                .mc-options {
                    margin: 15px 0;
                }
                
                .mc-option {
                    display: flex;
                    margin: 8px 0;
                    align-items: flex-start;
                }
                
                .option-checkbox {
                    width: 12px;
                    height: 12px;
                    border: 1.5px solid #2c3e50;
                    margin-right: 10px;
                    margin-top: 4px;
                    flex-shrink: 0;
                }
                
                .option-letter {
                    font-weight: 600;
                    color: #2c3e50;
                    margin-right: 8px;
                    min-width: 20px;
                }
                
                .option-text {
                    line-height: 1.5;
                    color: #34495e;
                }
                
                /* True/False Questions */
                .tf-options {
                    display: flex;
                    gap: 40px;
                    margin: 15px 0;
                    justify-content: flex-start;
                }
                
                .tf-option {
                    display: flex;
                    align-items: center;
                }
                
                /* Answer Spaces */
                .answer-space {
                    border-bottom: 2px solid #2c3e50;
                    display: inline-block;
                    min-width: 300px;
                    height: 25px;
                    margin: 10px 0;
                }
                
                .answer-lines {
                    margin: 15px 0;
                }
                
                .answer-line {
                    border-bottom: 1px solid #95a5a6;
                    height: 30px;
                    margin: 8px 0;
                    width: 100%;
                }
                
                /* Fill in the Blank */
                .fill-blank {
                    margin: 15px 0;
                }
                
                .blank-space {
                    border-bottom: 2px solid #2c3e50;
                    display: inline-block;
                    min-width: 120px;
                    height: 20px;
                    margin: 0 5px;
                }
                
                /* Essay Questions */
                .essay-space {
                    margin: 20px 0;
                    min-height: 200px;
                    border: 1px solid #bdc3c7;
                    background: repeating-linear-gradient(
                        transparent,
                        transparent 29px,
                        #ecf0f1 29px,
                        #ecf0f1 30px
                    );
                }
                
                /* Footer */
                .exam-footer {
                    margin-top: 50px;
                    padding-top: 20px;
                    border-top: 1px solid #bdc3c7;
                    text-align: center;
                    font-size: 12px;
                    color: #7f8c8d;
                }
                
                /* Print Styles */
                @media print {
                    body {
                        padding: 20px;
                        font-size: 12pt;
                    }
                    
                    .exam-container {
                        max-width: none;
                    }
                    
                    .question {
                        page-break-inside: avoid;
                        break-inside: avoid;
                    }
                    
                    .university-header {
                        page-break-after: avoid;
                    }
                    
                    .instructions-section {
                        page-break-after: avoid;
                    }
                }
                
                /* Instructor Answer Indicators */
                .correct-answer {
                    background-color: #d4edda !important;
                    border-color: #28a745 !important;
                }
                
                .correct-answer .option-checkbox {
                    background-color: #28a745;
                    border-color: #28a745;
                    position: relative;
                }
                
                .correct-answer .option-checkbox::after {
                    content: '✓';
                    color: white;
                    font-size: 10px;
                    position: absolute;
                    top: -2px;
                    left: 1px;
                }
                
                .answer-key {
                    margin-top: 10px;
                    padding: 8px 12px;
                    background-color: #e8f4fd;
                    border: 1px solid #3498db;
                    border-radius: 4px;
                    font-size: 13px;
                    color: #2980b9;
                }
                
                .answer-key strong {
                    color: #1565C0;
                }
                
                
                .option {
                    margin: 12px 0;
                    padding: 12px 15px;
                    background: #f9fafb;
                    border: 2px solid #e5e7eb;
                    border-radius: 10px;
                    display: flex;
                    align-items: center;
                    cursor: pointer;
                    transition: all 0.2s ease;
                }
            </style>
        </head>
        <body>
            <div class="exam-container">
                <!-- University Header -->
                <div class="university-header">
                    <div class="logo-placeholder">
                        LOGO
                    </div>
                    <div class="header-info">
                        <div class="university-name">{{ branding.university_name|default:"UNIVERSITY" }}</div>
                        <div class="department-name">{{ branding.department|default:"DEPARTMENT" }}</div>
                        <div class="course-info">{{ branding.course|default:"COURSE NAME" }} – {{ branding.semester|default:"SEMESTER 2025" }}</div>
                    </div>
                </div>
                
                <!-- Exam Title -->
                <div class="exam-title">
                    <div class="exam-type">{{ quiz_data.content_type|default:"EXAM"|upper }}</div>
                </div>
                
                <!-- Exam Metadata -->
                <div class="exam-metadata">
                    <div class="metadata-item">
                        <div class="metadata-label">INSTRUCTOR:</div>
                        <div class="metadata-value">{{ branding.instructor|default:"Full Name" }}</div>
                    </div>
                    <div class="metadata-item">
                        <div class="metadata-label">DATE:</div>
                        <div class="metadata-value">{{ branding.exam_date|default:"November 1, 2025" }}</div>
                    </div>
                    <div class="metadata-item">
                        <div class="metadata-label">DURATION:</div>
                        <div class="metadata-value">{{ quiz_data.estimated_duration|default:"2 hours" }}</div>
                    </div>
                    <div class="metadata-item">
                        <div class="metadata-label">TOTAL:</div>
                        <div class="metadata-value">{{ quiz_data.total_points|default:quiz_data.questions|length }} points</div>
                    </div>
                    <div class="metadata-item">
                        <div class="metadata-label">NAME:</div>
                        <div class="metadata-value">_________________________________</div>
                    </div>
                    <div class="metadata-item">
                        <div class="metadata-label">ID:</div>
                        <div class="metadata-value">_________________________________</div>
                    </div>
                </div>
                
                <!-- Instructions -->
                <div class="instructions-section">
                    <div class="instructions-title">Instructions:</div>
                    <ul class="instructions-list">
                        <li>Read each question carefully</li>
                        <li>Answer clearly in the space provided</li>
                        <li>Calculators are permitted unless otherwise noted</li>
                        <li>Show all work for calculation problems</li>
                    </ul>
                </div>
                
                <!-- Questions -->
                {{ questions_placeholder }}
                
                <!-- Footer -->
                <div class="exam-footer">
                    {{ branding.university_name|default:"University" }} – {{ branding.department|default:"Department" }} – Official {{ quiz_data.content_type|default:"Exam" }}
                </div>
            </div>
        </body>
        </html>
        """
        
        # Add instructor-specific CSS if showing answers
        if show_answers:
            instructor_css = """
            <style>
                /* Instructor version - correct answer styling */
                .correct-option {
                    background: #ecfdf5 !important;
                    border: 2px solid #10b981 !important;
                    font-weight: 600;
                }
                
                .correct-option .option-letter {
                    background: #10b981;
                    color: white;
                    padding: 2px 8px;
                    border-radius: 4px;
                }
                
                /* Expected answer styling for instructor version */
                .expected-answer {
                    background: #f0f9ff;
                    border: 1px solid #0ea5e9;
                    border-radius: 6px;
                    padding: 8px 12px;
                    margin-top: 10px;
                    color: #0c4a6e;
                    font-size: 0.9em;
                }
                
                .expected-answer strong {
                    color: #0369a1;
                }
            </style>
            """
            # Insert the instructor CSS before </head>
            html_template = html_template.replace('</head>', instructor_css + '\n        </head>')
        
        # Simple template rendering (replace all placeholders)
        html = html_template
        
        # Basic replacements
        html = html.replace('{{ quiz_data.title }}', quiz_data.get('title', 'Quiz'))
        html = html.replace('{{ quiz_data.description }}', quiz_data.get('description', ''))
        html = html.replace('{{ export_date }}', datetime.now().strftime('%B %d, %Y'))
        
        # Branding replacements
        html = html.replace('{{ branding.university_name|default:"UNIVERSITY" }}', branding.get('university_name', 'UNIVERSITY') if branding else 'UNIVERSITY')
        html = html.replace('{{ branding.department|default:"DEPARTMENT" }}', branding.get('department', 'DEPARTMENT') if branding else 'DEPARTMENT')
        html = html.replace('{{ branding.course|default:"COURSE NAME" }}', branding.get('course', 'COURSE NAME') if branding else 'COURSE NAME')
        html = html.replace('{{ branding.semester|default:"SEMESTER 2025" }}', branding.get('semester', 'SEMESTER 2025') if branding else 'SEMESTER 2025')
        html = html.replace('{{ branding.instructor|default:"Full Name" }}', branding.get('instructor', 'Full Name') if branding else 'Full Name')
        html = html.replace('{{ branding.exam_date|default:"November 1, 2025" }}', branding.get('exam_date', 'November 1, 2025') if branding else 'November 1, 2025')
        
        # More comprehensive template cleanup
        import re
        
        # Replace branding placeholders
        if branding and branding.get('institution_name'):
            html = re.sub(r'{% if branding\.institution_name %}\s*<div class="institution">{{ branding\.institution_name }}</div>\s*{% endif %}', 
                         f'<div class="institution">{branding["institution_name"]}</div>', html, flags=re.DOTALL)
        else:
            html = re.sub(r'{% if branding\.institution_name %}.*?{% endif %}', '', html, flags=re.DOTALL)
            
        if branding and branding.get('department'):
            html = re.sub(r'{% if branding\.department %}\s*<div class="department">{{ branding\.department }}</div>\s*{% endif %}', 
                         f'<div class="department">{branding["department"]}</div>', html, flags=re.DOTALL)
        else:
            html = re.sub(r'{% if branding\.department %}.*?{% endif %}', '', html, flags=re.DOTALL)
        
        # Replace description placeholder more comprehensively
        description_placeholder = '''{% if quiz_data.description %}
                    <p class="description">{{ quiz_data.description }}</p>
                    {% endif %}'''
        if quiz_data.get('description'):
            html = html.replace(description_placeholder, f'<p class="description">{quiz_data["description"]}</p>')
        else:
            html = html.replace(description_placeholder, '')
        
        # Replace quiz data placeholders
        html = html.replace('{{ quiz_data.estimated_duration|default:"Not specified" }}', 
                           quiz_data.get('estimated_duration', 'Not specified'))
        html = html.replace('{{ quiz_data.total_points|default:quiz_data.questions|length }}', 
                           str(quiz_data.get('total_points', len(quiz_data.get('questions', [])))))
        
        # Add questions with proper type distinction
        questions_html = ""
        for i, question in enumerate(quiz_data.get('questions', []), 1):
            question_type = question.get('type', 'multiple_choice')
            question_type_display = {
                'multiple_choice': 'Çoktan Seçmeli' if 'tr' in str(quiz_data.get('language', '')).lower() else 'Multiple Choice',
                'true_false': 'Doğru/Yanlış' if 'tr' in str(quiz_data.get('language', '')).lower() else 'True/False',
                'short_answer': 'Kısa Cevap' if 'tr' in str(quiz_data.get('language', '')).lower() else 'Short Answer',
                'fill_blank': 'Boşluk Doldurma' if 'tr' in str(quiz_data.get('language', '')).lower() else 'Fill in the Blank',
                'essay': 'Kompozisyon' if 'tr' in str(quiz_data.get('language', '')).lower() else 'Essay'
            }.get(question_type, 'Multiple Choice')
            
            # Escape HTML special characters in question content
            import html as html_escape_module
            question_text = html_escape_module.escape(question.get('question', '')) if question.get('question') else ''
            
            # University-style question formatting
            q_html = f"""
                <div class="question">
                    <div class="question-header">
                        <div class="question-number">{i}.</div>
                        <div class="question-points">({question.get('points', 1)} point{'s' if question.get('points', 1) != 1 else ''})</div>
                    </div>
                    <div class="question-text">{question_text}</div>
            """
            
            # Handle different question types
            if question_type == 'multiple_choice':
                if question.get('options') and len(question['options']) > 0:
                    correct_answer = question.get('correct_answer', '').upper() if show_answers else ''
                    q_html += '<div class="mc-options">'
                    for j, option in enumerate(question['options']):
                        option_letter = chr(65 + j)  # A, B, C, D...
                        escaped_option = html_escape_module.escape(str(option)) if option else ''
                        
                        # Only show correct answer indicators if show_answers is True
                        correct_class = ' correct-answer' if show_answers and correct_answer == option_letter else ''
                        
                        q_html += f'''
                            <div class="mc-option{correct_class}">
                                <div class="option-checkbox"></div>
                                <div class="option-letter">{option_letter}.</div>
                                <div class="option-text">{escaped_option}</div>
                            </div>
                        '''
                    q_html += '</div>'
                    
                    # Add explanation for instructor version
                    if show_answers and question.get('explanation'):
                        explanation = html_escape_module.escape(str(question.get('explanation', '')))
                        q_html += f'<div class="answer-key"><strong>Explanation:</strong> {explanation}</div>'
                else:
                    # Fallback if no options provided
                    q_html += '<div class="mc-options">'
                    for j, letter in enumerate(['A', 'B', 'C', 'D']):
                        q_html += f'''
                            <div class="mc-option">
                                <div class="option-checkbox"></div>
                                <div class="option-letter">{letter}.</div>
                                <div class="option-text">___________</div>
                            </div>
                        '''
                    q_html += '</div>'
            
            elif question_type == 'true_false':
                true_label = 'True'
                false_label = 'False'
                correct_answer = str(question.get('correct_answer', '')).lower() if show_answers else ''
                
                true_class = ' correct-answer' if show_answers and correct_answer == 'true' else ''
                false_class = ' correct-answer' if show_answers and correct_answer == 'false' else ''
                
                q_html += f'''
                    <div class="tf-options">
                        <div class="tf-option{true_class}">
                            <div class="option-checkbox"></div>
                            <span>{true_label}</span>
                        </div>
                        <div class="tf-option{false_class}">
                            <div class="option-checkbox"></div>
                            <span>{false_label}</span>
                        </div>
                    </div>
                '''
            
            elif question_type == 'short_answer':
                q_html += '<div class="answer-space"></div>'
                
                # Show expected answer and explanation for instructor version
                if show_answers and question.get('correct_answer'):
                    expected_answer = html_escape_module.escape(str(question['correct_answer']))
                    explanation = html_escape_module.escape(str(question.get('explanation', '')))
                    q_html += f'<div class="answer-key"><strong>Expected Answer:</strong> {expected_answer}'
                    if explanation:
                        q_html += f'<br><strong>Explanation:</strong> {explanation}'
                    q_html += '</div>'
            
            elif question_type == 'fill_blank':
                # Add blank spaces for fill in the blank questions
                q_html += '<div class="fill-blank">'
                q_html += '<div class="blank-space"></div>'
                q_html += '</div>'
                
                # Show expected answer for instructor version
                if show_answers and question.get('correct_answer'):
                    expected_answer = html_escape_module.escape(str(question['correct_answer']))
                    q_html += f'<div class="answer-key"><strong>Expected Answer:</strong> {expected_answer}</div>'
            
            elif question_type == 'essay':
                # Use the essay-space styling for lined writing area
                q_html += '<div class="essay-space"></div>'
                
                # Show expected answer for instructor version
                if show_answers and question.get('correct_answer'):
                    expected_answer = html_escape_module.escape(str(question['correct_answer']))
                    q_html += f'<div class="answer-key"><strong>Expected Answer:</strong> {expected_answer}</div>'
            
            # Add catch-all for any unhandled question type
            else:
                answer_label = 'Cevap:' if 'tr' in str(quiz_data.get('language', '')).lower() else 'Answer:'
                answer_space = '<span class="answer-space"></span>'
                
                # Show expected answer for instructor version
                if show_answers and question.get('correct_answer'):
                    expected_answer = html_escape_module.escape(str(question['correct_answer']))
                    answer_space += f'<div class="expected-answer"><strong>Expected:</strong> {expected_answer}</div>'
                
                q_html += f'<div class="answer-container">{answer_label} {answer_space}</div>'
            
            q_html += '</div>'
            questions_html += q_html
        
        # Replace quiz data placeholders for content type and other fields
        html = html.replace('{{ quiz_data.content_type|default:"EXAM"|upper }}', quiz_data.get('content_type', 'QUIZ').upper())
        html = html.replace('{{ quiz_data.estimated_duration|default:"2 hours" }}', quiz_data.get('estimated_duration', '2 hours'))
        html = html.replace('{{ quiz_data.total_points|default:quiz_data.questions|length }} points', f"{quiz_data.get('total_points', len(quiz_data.get('questions', [])))} points")
        
        # Footer replacements
        html = html.replace('{{ branding.university_name|default:"University" }}', branding.get('university_name', 'University') if branding else 'University')
        html = html.replace('{{ quiz_data.content_type|default:"Exam" }}', quiz_data.get('content_type', 'Quiz'))
        
        # Replace questions placeholder with generated content
        html = html.replace('{{ questions_placeholder }}', questions_html)
        
        return html


class ZIPExporter:
    """Service for creating ZIP archives with multiple formats"""
    
    def __init__(self):
        self.pdf_exporter = PDFExporter()
        self.html_exporter = HTMLExporter()
        if DOCX_AVAILABLE:
            self.docx_exporter = DOCXExporter()
    
    def export_complete_package(self, content_data: Dict[str, Any], 
                              versions: List[str] = None,
                              formats: List[str] = None,
                              branding: Dict[str, Any] = None) -> io.BytesIO:
        """
        Create a complete export package with multiple formats and versions
        
        Args:
            content_data: The quiz/exam data
            versions: List of version letters (e.g., ['A', 'B', 'C'])
            formats: List of formats to include ('pdf', 'docx', 'html')
            branding: Branding information
            
        Returns:
            BytesIO buffer containing the ZIP file
        """
        buffer = io.BytesIO()
        versions = versions or ['A']
        formats = formats or ['pdf']
        
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            content_type = content_data.get('type', 'quiz')
            base_title = content_data.get('title', 'Content')
            
            for version in versions:
                # Create version-specific data
                version_data = self._create_version_data(content_data, version)
                
                for format_type in formats:
                    try:
                        if format_type == 'pdf':
                            # Export questions
                            if content_type == 'quiz':
                                pdf_buffer = self.pdf_exporter.export_quiz(version_data, branding)
                            else:
                                pdf_buffer = self.pdf_exporter.export_exam(version_data, branding)
                            
                            zipf.writestr(f'{base_title}_Version_{version}.pdf', pdf_buffer.getvalue())
                            
                            # Export answer key
                            answer_key_buffer = self.pdf_exporter.export_answer_key(version_data, branding)
                            zipf.writestr(f'{base_title}_Version_{version}_Answer_Key.pdf', answer_key_buffer.getvalue())
                        
                        elif format_type == 'docx' and DOCX_AVAILABLE:
                            if content_type == 'quiz':
                                docx_buffer = self.docx_exporter.export_quiz(version_data, branding)
                            else:
                                docx_buffer = self.docx_exporter.export_exam(version_data, branding)
                            
                            zipf.writestr(f'{base_title}_Version_{version}.docx', docx_buffer.getvalue())
                        
                        elif format_type == 'html':
                            html_content = self.html_exporter.export_quiz(version_data, branding)
                            zipf.writestr(f'{base_title}_Version_{version}.html', html_content.encode('utf-8'))
                    
                    except Exception as e:
                        logger.error(f"Error exporting {format_type} for version {version}: {str(e)}")
                        continue
            
            # Add metadata file
            metadata = {
                'title': base_title,
                'content_type': content_type,
                'versions': versions,
                'formats': formats,
                'export_date': datetime.now().isoformat(),
                'branding': branding or {}
            }
            zipf.writestr('metadata.json', json.dumps(metadata, indent=2))
        
        buffer.seek(0)
        return buffer
    
    def _create_version_data(self, content_data: Dict[str, Any], version: str) -> Dict[str, Any]:
        """Create version-specific content data"""
        import random
        
        version_data = content_data.copy()
        version_data['title'] = f"{content_data.get('title', 'Content')} - Version {version}"
        
        # Shuffle questions for different versions
        if 'questions' in version_data:
            questions = version_data['questions'].copy()
            random.seed(ord(version))  # Use version letter as seed for reproducibility
            random.shuffle(questions)
            version_data['questions'] = questions
        elif 'sections' in version_data:
            # Handle exam format
            for section in version_data['sections']:
                if 'questions' in section:
                    questions = section['questions'].copy()
                    random.seed(ord(version))
                    random.shuffle(questions)
                    section['questions'] = questions
        
        return version_data


class ExportService:
    """Main export service that coordinates all export types"""
    
    def __init__(self):
        self.pdf_exporter = PDFExporter()
        self.html_exporter = HTMLExporter()
        self.zip_exporter = ZIPExporter()
        
        if DOCX_AVAILABLE:
            self.docx_exporter = DOCXExporter()
    
    def export_content(self, content_data: Dict[str, Any], 
                      export_format: str,
                      branding: Dict[str, Any] = None,
                      include_answer_key: bool = True,
                      versions: List[str] = None) -> Dict[str, Any]:
        """
        Main export method that handles all formats
        
        Args:
            content_data: The content to export
            export_format: 'pdf', 'docx', 'html', or 'zip'
            branding: Branding information
            include_answer_key: Whether to include answer key
            versions: Version letters for multi-version exports
            
        Returns:
            Dict with export results
        """
        try:
            if export_format == 'pdf':
                return self._export_pdf(content_data, branding, include_answer_key)
            elif export_format == 'docx':
                return self._export_docx(content_data, branding)
            elif export_format == 'html':
                return self._export_html(content_data, branding)
            elif export_format == 'json':
                return self._export_json(content_data, branding)
            elif export_format == 'zip':
                return self._export_zip(content_data, branding, versions or ['A', 'B', 'C'])
            else:
                return {
                    'success': False,
                    'error': f'Unsupported export format: {export_format}'
                }
        
        except Exception as e:
            logger.error(f"Export error: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _export_pdf(self, content_data: Dict[str, Any], branding: Dict[str, Any], include_answer_key: bool) -> Dict[str, Any]:
        """Export to PDF format"""
        content_type = content_data.get('type', 'quiz')
        
        if content_type == 'quiz':
            buffer = self.pdf_exporter.export_quiz(content_data, branding)
        else:
            buffer = self.pdf_exporter.export_exam(content_data, branding)
        
        result = {
            'success': True,
            'file_data': buffer.getvalue(),
            'filename': f"{content_data.get('title', 'content')}.pdf",
            'content_type': 'application/pdf'
        }
        
        if include_answer_key:
            answer_key_buffer = self.pdf_exporter.export_answer_key(content_data, branding)
            result['answer_key_data'] = answer_key_buffer.getvalue()
            result['answer_key_filename'] = f"{content_data.get('title', 'content')}_answer_key.pdf"
        
        return result
    
    def _export_docx(self, content_data: Dict[str, Any], branding: Dict[str, Any]) -> Dict[str, Any]:
        """Export to DOCX format"""
        if not DOCX_AVAILABLE:
            return {
                'success': False,
                'error': 'DOCX export not available. Install python-docx.'
            }
        
        content_type = content_data.get('type', 'quiz')
        
        if content_type == 'quiz':
            buffer = self.docx_exporter.export_quiz(content_data, branding)
        else:
            buffer = self.docx_exporter.export_exam(content_data, branding)
        
        return {
            'success': True,
            'file_data': buffer.getvalue(),
            'filename': f"{content_data.get('title', 'content')}.docx",
            'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
    
    def _export_html(self, content_data: Dict[str, Any], branding: Dict[str, Any]) -> Dict[str, Any]:
        """Export to HTML format"""
        html_content = self.html_exporter.export_quiz(content_data, branding)
        
        return {
            'success': True,
            'file_data': html_content.encode('utf-8'),
            'filename': f"{content_data.get('title', 'content')}.html",
            'content_type': 'text/html'
        }
    
    def _export_json(self, content_data: Dict[str, Any], branding: Dict[str, Any]) -> Dict[str, Any]:
        """Export to JSON format"""
        # Add metadata to content
        export_data = {
            'export_metadata': {
                'export_date': datetime.now().isoformat(),
                'format': 'json',
                'branding': branding
            },
            'content': content_data
        }
        
        json_content = json.dumps(export_data, indent=2, default=str)
        
        return {
            'success': True,
            'file_data': json_content.encode('utf-8'),
            'filename': f"{content_data.get('title', 'content')}.json",
            'content_type': 'application/json'
        }
    
    def _export_zip(self, content_data: Dict[str, Any], branding: Dict[str, Any], versions: List[str]) -> Dict[str, Any]:
        """Export to ZIP format with multiple versions"""
        buffer = self.zip_exporter.export_complete_package(
            content_data, 
            versions=versions,
            formats=['pdf', 'html'] + (['docx'] if DOCX_AVAILABLE else []),
            branding=branding
        )
        
        return {
            'success': True,
            'file_data': buffer.getvalue(),
            'filename': f"{content_data.get('title', 'content')}_complete_package.zip",
            'content_type': 'application/zip'
        }
    
    def export_generation(self, export_job) -> Dict[str, Any]:
        """
        Export an AI generation to the specified format
        
        Args:
            export_job: ExportJob instance containing export configuration
            
        Returns:
            Dict with export results
        """
        try:
            # Get generation data
            generation = export_job.generation
            
            # Prepare content data
            content_data = self._prepare_generation_data(generation)
            
            # Prepare branding data
            branding = export_job.branding_settings or {}
            if export_job.watermark:
                branding['watermark'] = export_job.watermark
            
            # Export using the main export method
            result = self.export_content(
                content_data=content_data,
                export_format=export_job.export_format,
                branding=branding,
                include_answer_key=export_job.include_answer_key
            )
            
            if result['success']:
                # Save the exported file
                file_content = ContentFile(result['file_data'])
                export_job.generated_file.save(
                    result['filename'],
                    file_content,
                    save=False  # Don't save yet, we need to set file_size first
                )
                export_job.file_size = len(result['file_data'])
                export_job.save()  # Save the export_job with file_size
                export_job.mark_completed()
                
                # Save answer key if available
                if 'answer_key_data' in result:
                    # Create a separate export for answer key or store as metadata
                    export_job.parameters['answer_key_available'] = True
                    export_job.save()
            else:
                export_job.mark_error(result.get('error', 'Unknown export error'))
            
            return result
            
        except Exception as e:
            error_msg = f"Export generation failed: {str(e)}"
            logger.error(error_msg)
            export_job.mark_error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }
    
    def export_with_versions(self, export_job, num_versions: int = 3) -> Dict[str, Any]:
        """
        Export an AI generation with multiple versions (A, B, C)
        
        Args:
            export_job: ExportJob instance containing export configuration
            num_versions: Number of versions to create
            
        Returns:
            Dict with export results
        """
        try:
            # Get generation data
            generation = export_job.generation
            
            # Prepare content data
            content_data = self._prepare_generation_data(generation)
            
            # Prepare branding data
            branding = export_job.branding_settings or {}
            if export_job.watermark:
                branding['watermark'] = export_job.watermark
            
            # Create version letters
            version_letters = [chr(65 + i) for i in range(num_versions)]  # A, B, C, etc.
            
            if export_job.export_format == 'zip':
                # Export as ZIP with multiple versions
                result = self._export_zip(content_data, branding, version_letters)
            else:
                # For non-ZIP formats, create individual version files
                result = self._export_individual_versions(
                    export_job, content_data, branding, version_letters
                )
            
            if result['success']:
                # Save the main export file
                file_content = ContentFile(result['file_data'])
                export_job.generated_file.save(
                    result['filename'],
                    file_content,
                    save=False
                )
                export_job.file_size = len(result['file_data'])
                export_job.save()  # Save the export_job with file_size
                export_job.mark_completed()
            else:
                export_job.mark_error(result.get('error', 'Unknown export error'))
            
            return result
            
        except Exception as e:
            error_msg = f"Export with versions failed: {str(e)}"
            logger.error(error_msg)
            export_job.mark_error(error_msg)
            return {
                'success': False,
                'error': error_msg
            }
    
    def _prepare_generation_data(self, generation) -> Dict[str, Any]:
        """
        Convert AI generation to export-ready data format
        
        Args:
            generation: AIGeneration instance
            
        Returns:
            Dict containing structured content data
        """
        content_data = {
            'title': generation.title,
            'description': generation.description or '',
            'type': generation.content_type,
            'created_at': generation.created_at,
        }
        
        # Handle generated content
        if generation.generated_content:
            generated = generation.generated_content
            
            # Extract questions if available
            if 'questions' in generated:
                content_data['questions'] = generated['questions']
            elif 'sections' in generated:
                content_data['sections'] = generated['sections']
            
            # Extract metadata
            content_data.update({
                'total_points': generated.get('total_points', 0),
                'estimated_duration': generated.get('estimated_duration', ''),
            })
        
        # Handle individual questions from database
        if hasattr(generation, 'questions') and generation.questions.exists():
            questions = []
            for q in generation.questions.all():
                question_data = {
                    'id': q.id,
                    'type': q.question_type,
                    'question': q.question_text,
                    'points': q.points,
                    'difficulty': q.difficulty,
                    'correct_answer': q.correct_answer,
                    'explanation': q.explanation or '',
                }
                
                # Add options for multiple choice questions
                if q.question_type == 'mcq' and q.options:
                    question_data['options'] = q.options
                
                questions.append(question_data)
            
            content_data['questions'] = questions
            content_data['total_points'] = sum(q.points for q in generation.questions.all())
        
        return content_data
    
    def _export_individual_versions(self, export_job, content_data: Dict[str, Any], 
                                   branding: Dict[str, Any], version_letters: List[str]) -> Dict[str, Any]:
        """
        Export individual version files and create ExportVersion records
        
        Args:
            export_job: ExportJob instance
            content_data: Content to export
            branding: Branding information
            version_letters: List of version letters
            
        Returns:
            Dict with export results (returns the first version as main file)
        """
        from .models import ExportVersion
        import random
        
        results = []
        main_result = None
        
        for i, version_letter in enumerate(version_letters):
            # Create version-specific data
            version_data = content_data.copy()
            version_data['title'] = f"{content_data.get('title', 'Content')} - Version {version_letter}"
            
            # Randomize question order for versions
            if 'questions' in version_data:
                questions = version_data['questions'].copy()
                random.seed(ord(version_letter))  # Reproducible randomization
                random.shuffle(questions)
                version_data['questions'] = questions
            
            # Export this version
            result = self.export_content(
                content_data=version_data,
                export_format=export_job.export_format,
                branding=branding,
                include_answer_key=export_job.include_answer_key
            )
            
            if result['success']:
                # Create ExportVersion record
                export_version = ExportVersion.objects.create(
                    export_job=export_job,
                    version_letter=version_letter,
                    file_size=len(result['file_data']),
                    variations={'randomized_order': True}
                )
                
                # Save version file
                version_filename = f"{content_data.get('title', 'content')}_Version_{version_letter}.{export_job.export_format}"
                file_content = ContentFile(result['file_data'])
                export_version.generated_file.save(
                    version_filename,
                    file_content,
                    save=True
                )
                
                # Use first version as main result
                if i == 0:
                    main_result = result
                    main_result['filename'] = version_filename
                
                results.append(result)
        
        return main_result or {
            'success': False,
            'error': 'Failed to create any versions'
        }
